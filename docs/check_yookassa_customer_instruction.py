from pathlib import Path
from zipfile import ZipFile

from docx import Document


path = next((Path.cwd() / "docs").glob("*Arasaka_Buisness.docx"))
doc = Document(path)

texts = []
for paragraph in doc.paragraphs:
    if paragraph.text.strip():
        texts.append(paragraph.text.strip())

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if paragraph.text.strip():
                    texts.append(paragraph.text.strip())

required = [
    "Инструкция для подключения ЮKassa",
    "Что отправить одним сообщением",
    "Тестовый shopId",
    "Настоящий Secret key",
    "HTTP-уведомления",
    "payment.succeeded",
    "payment.canceled",
    "Чеки и бухгалтерия",
    "пополнение внутреннего кошелька",
    "Arasaka Buisness",
]

missing = [item for item in required if not any(item in text for text in texts)]

with ZipFile(path) as archive:
    names = set(archive.namelist())
    document_xml = archive.read("word/document.xml").decode("utf-8")
    numbering_xml = archive.read("word/numbering.xml").decode("utf-8") if "word/numbering.xml" in names else ""

print(f"path={path}")
print(f"size={path.stat().st_size}")
print(f"paragraphs={len(doc.paragraphs)}")
print(f"tables={len(doc.tables)}")
print(f"required_missing={missing}")
print(f"has_numbering_xml={bool(numbering_xml)}")
print(f"numbering_count={numbering_xml.count('<w:num ')}")
print(f"table_width_tokens={document_xml.count('w:tblW')}")
print(f"api_webhook_count={document_xml.count('api/v1/payments/yookassa/webhook')}")
