from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Инструкция_ЮKassa_для_заказчика_Arasaka_Buisness_обновлено.docx"

BRAND = "Arasaka Buisness"
PROJECT = "Портал «Воевода»"
API_BASE_URL = "https://api.voevoda.ru"
WEBHOOK_URL = f"{API_BASE_URL}/api/v1/payments/yookassa/webhook/"
ACCENT = "1F4D78"
ACCENT_LIGHT = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
MID_GRAY = "D9E2EC"
TEXT = "1F2933"
MUTED = "667085"
CONTENT_DXA = 9360


def set_font(run, name="Calibri", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_font(run, size=9, color=MUTED)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color="D0D7DE"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        border = borders.find(tag)
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:w"), str(indent_dxa))
    ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_numbering_definition(doc, kind="decimal"):
    numbering = doc.part.numbering_part.element
    abstract_id = str(len(numbering.findall(qn("w:abstractNum"))) + 20)
    num_id = str(len(numbering.findall(qn("w:num"))) + 20)

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), abstract_id)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if kind == "decimal" else "bullet")
    lvl.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if kind == "decimal" else "•")
    lvl.append(lvl_text)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(tabs)
    p_pr.append(ind)
    lvl.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:cs"), "Calibri")
    r_pr.append(r_fonts)
    lvl.append(r_pr)

    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    num.append(abstract_ref)
    numbering.append(num)
    return int(num_id)


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    set_paragraph_spacing(paragraph, after=4, line=1.25)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Calibri")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:cs"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True


def add_text(paragraph, text, bold=False, color=TEXT, size=11):
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return run


def add_h1(doc, text):
    p = doc.add_paragraph(text, style="Heading 1")
    return p


def add_h2(doc, text):
    p = doc.add_paragraph(text, style="Heading 2")
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, bold=True)
        add_text(p, text[len(bold_prefix):])
    else:
        add_text(p, text)
    return p


def add_bullet(doc, text, bullet_num_id):
    p = doc.add_paragraph()
    apply_numbering(p, bullet_num_id)
    add_text(p, text)
    return p


def add_step(doc, text, decimal_num_id):
    p = doc.add_paragraph()
    apply_numbering(p, decimal_num_id)
    add_text(p, text)
    return p


def add_callout(doc, title, body, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA], indent_dxa=120)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(3)
    p = cell.paragraphs[0]
    add_text(p, title, bold=True, color=ACCENT, size=10.5)
    p2 = cell.add_paragraph()
    set_paragraph_spacing(p2, after=0, line=1.2)
    add_text(p2, body, size=10.5)
    return table


def add_label_table(doc, rows, widths=(2700, 6660), header=None):
    table = doc.add_table(rows=0, cols=2)
    if header:
        row = table.add_row()
        row.cells[0].merge(row.cells[1])
        shade_cell(row.cells[0], ACCENT_LIGHT)
        p = row.cells[0].paragraphs[0]
        set_paragraph_spacing(p, after=0, line=1.2)
        add_text(p, header, bold=True, color=ACCENT)
    for label, detail in rows:
        row = table.add_row()
        shade_cell(row.cells[0], LIGHT_GRAY)
        p0 = row.cells[0].paragraphs[0]
        set_paragraph_spacing(p0, after=0, line=1.2)
        add_text(p0, label, bold=True, size=10.5)
        p1 = row.cells[1].paragraphs[0]
        set_paragraph_spacing(p1, after=0, line=1.2)
        add_text(p1, detail, size=10.5)
    set_table_geometry(table, list(widths), indent_dxa=120)
    return table


def add_checklist_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    widths = [900, 3420, 5040]
    headers = ["Готово", "Пункт", "Что должно быть сделано"]
    for idx, text in enumerate(headers):
        shade_cell(table.cell(0, idx), ACCENT_LIGHT)
        p = table.cell(0, idx).paragraphs[0]
        set_paragraph_spacing(p, after=0, line=1.15)
        add_text(p, text, bold=True, color=ACCENT, size=10)
    for mark, item, detail in rows:
        row = table.add_row()
        for idx, value in enumerate((mark, item, detail)):
            p = row.cells[idx].paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.15)
            add_text(p, value, size=10.2)
    set_table_geometry(table, widths, indent_dxa=120)
    return table


def add_copy_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA], indent_dxa=120)
    cell = table.cell(0, 0)
    shade_cell(cell, "FAFBFC")
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line=1.15)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_font(run, name="Consolas", size=9.2, color=TEXT)
    return table


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)
    decimal_num_id = add_numbering_definition(doc, "decimal")
    bullet_num_id = add_numbering_definition(doc, "bullet")

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(header_p, BRAND, bold=True, color=ACCENT, size=9.5)
    add_text(header_p, " | инструкция для заказчика", color=MUTED, size=9.5)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    add_page_number(footer_p)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.line_spacing = 1.1
    add_text(title, "Инструкция для подключения ЮKassa", bold=True, color=ACCENT, size=22)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    add_text(subtitle, "Что нужно сделать заказчику и какие данные передать ", color=MUTED, size=11)
    add_text(subtitle, BRAND, bold=True, color=ACCENT, size=11)

    add_label_table(doc, [
        ("Проект", PROJECT),
        ("Формат оплаты", "Переход пользователя на защищенную платежную страницу ЮKassa."),
        ("Стоимость варианта", "На старте дополнительные платные модули покупать не нужно. Если что-то потребуется, Arasaka Buisness отдельно сообщит, что именно купить и зачем."),
        ("Где вводится карта", "Только на стороне ЮKassa. Сайт не хранит и не обрабатывает данные карт."),
        ("API-домен", API_BASE_URL),
        ("Дата документа", "05.06.2026"),
        ("Версия", "1.0"),
    ], widths=(2100, 7260), header="Краткая информация")

    add_callout(
        doc,
        "Главная задача",
        "Пройти пункты ниже по порядку и прислать Arasaka Buisness данные из раздела «Что отправить одним сообщением». "
        "Если какого-то пункта нет в кабинете ЮKassa, сделайте скриншот этого места и отправьте нам.",
        fill="F4F6F9",
    )

    add_h1(doc, "1. Что именно подключаем")
    for text in [
        "Оплату курсов и товаров через ЮKassa.",
        "Пополнение внутреннего кошелька пользователя через ЮKassa.",
        "Историю платежей в личном кабинете пользователя.",
        "Проверку платежа через уведомления ЮKassa, чтобы сайт отмечал оплату только после реального успешного платежа.",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_h1(doc, "2. Что подготовить до настройки")
    for text in [
        "Доступ в личный кабинет ЮKassa с правами владельца или администратора.",
        "Возможность создать тестовый магазин и настоящий магазин.",
        "Адрес сайта, на котором будет работать портал.",
        f"Адрес API-сервера для ЮKassa: {API_BASE_URL}.",
        "Ответ бухгалтера по онлайн-чекам: нужны ли чеки через ЮKassa, какая система налогообложения и какая ставка НДС.",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_h1(doc, "3. Что сделать в ЮKassa: тестовый магазин")
    for text in [
        "Войдите в личный кабинет ЮKassa.",
        "Нажмите на название текущего магазина в верхней части кабинета.",
        "Нажмите «Добавить магазин».",
        "Выберите тестовый магазин.",
        "Выберите прием платежей на сайте.",
        "Укажите адрес сайта. Если боевого домена пока нет, укажите текущий тестовый адрес или уточните его у Arasaka Buisness.",
        "Сохраните магазин и откройте его настройки.",
    ]:
        add_step(doc, text, decimal_num_id)

    add_h2(doc, "Какие данные взять из тестового магазина")
    add_label_table(doc, [
        ("shopId", "Идентификатор магазина. Он нужен серверу, чтобы создавать платежи."),
        ("Secret key", "Секретный API-ключ магазина. Он нужен только серверу. Его нельзя публиковать и нельзя вставлять в сайт."),
        ("Скриншот", "Необязательно, но удобно: скриншот страницы, где видно название тестового магазина и shopId."),
    ], widths=(2200, 7160))

    add_h1(doc, "4. Где найти shopId")
    for text in [
        "Откройте нужный магазин в кабинете ЮKassa.",
        "Перейдите в раздел «Настройки».",
        "Откройте раздел «Магазин».",
        "Найдите поле shopId.",
        "Скопируйте значение и отправьте его Arasaka Buisness.",
    ]:
        add_step(doc, text, decimal_num_id)

    add_h1(doc, "5. Где найти Secret key")
    for text in [
        "Откройте нужный магазин в кабинете ЮKassa.",
        "Перейдите в раздел «Интеграция».",
        "Откройте «Ключи API».",
        "Если ключ уже создан и виден, скопируйте его.",
        "Если ключ не виден, нажмите «Выпустить ключ» или «Перевыпустить ключ».",
        "Скопируйте ключ сразу после создания и отправьте его только в личном безопасном сообщении.",
    ]:
        add_step(doc, text, decimal_num_id)

    add_callout(
        doc,
        "Важно про Secret key",
        "Secret key нельзя отправлять в общий чат, публиковать в документах с общим доступом или вставлять во фронтенд. "
        "Его нужно передать только ответственному разработчику Arasaka Buisness.",
        fill="FFF7E6",
    )

    add_h1(doc, "6. Настройка HTTP-уведомлений")
    add_body(
        doc,
        "HTTP-уведомления нужны, чтобы сайт автоматически узнавал, что платеж прошел успешно или был отменен.",
    )
    for text in [
        "Откройте нужный магазин в кабинете ЮKassa.",
        "Перейдите в раздел «Интеграция».",
        "Откройте «HTTP-уведомления».",
        "Нажмите «Изменить настройки».",
        "В поле URL укажите адрес уведомлений.",
        "Включите события payment.succeeded и payment.canceled.",
        "Сохраните настройки.",
    ]:
        add_step(doc, text, decimal_num_id)

    add_label_table(doc, [
        ("URL для уведомлений", WEBHOOK_URL),
        ("Что заменить", "Ничего заменять не нужно. Этот URL уже готов для передачи заказчику."),
        ("События", "payment.succeeded и payment.canceled"),
    ], widths=(2450, 6910), header="Параметры уведомлений")

    add_h1(doc, "7. Чеки и бухгалтерия")
    add_body(
        doc,
        "Этот блок нужно согласовать с бухгалтером заказчика до включения чеков. Без подтверждения бухгалтера чеки через ЮKassa включать не нужно.",
    )
    for text in [
        "Нужны ли онлайн-чеки именно через ЮKassa: да или нет.",
        "Какая система налогообложения используется.",
        "Какая ставка НДС должна быть в чеке. Точную ставку должен подтвердить бухгалтер.",
        "Как проводить курсы в чеке: услуга, предоплата, полная оплата или другой вариант бухгалтера.",
        "Как проводить товары в чеке: товар, нужна ли маркировка, есть ли особые требования.",
        "Как проводить пополнение внутреннего кошелька: аванс, предоплата или другой вариант бухгалтера.",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_callout(
        doc,
        "Что просить у заказчика по чекам",
        "Попросите заказчика написать один короткий ответ от бухгалтера: «Чеки через ЮKassa нужны/не нужны. "
        "Система налогообложения: ____. НДС: ____. Курсы: ____. Товары: ____. Пополнение кошелька: ____».",
        fill=ACCENT_LIGHT,
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    section2 = doc.sections[-1]
    section2.header_distance = Inches(0.492)
    section2.footer_distance = Inches(0.492)

    add_h1(doc, "8. После теста: настоящий магазин")
    add_body(
        doc,
        "Когда тестовый платеж пройдет успешно, нужно повторить те же действия для настоящего магазина.",
    )
    for text in [
        "Откройте настоящий магазин в ЮKassa.",
        "Скопируйте production shopId.",
        "Скопируйте или выпустите production Secret key.",
        f"Настройте HTTP-уведомления на URL: {WEBHOOK_URL}.",
        "Включите события payment.succeeded и payment.canceled.",
        "Проверьте, какие способы оплаты реально подключены в договоре: банковские карты, СБП, ЮMoney и другие.",
        "Передайте production shopId и production Secret key Arasaka Buisness.",
    ]:
        add_step(doc, text, decimal_num_id)

    add_h1(doc, "9. Что отправить одним сообщением")
    add_body(
        doc,
        "Ниже шаблон. Его можно скопировать, заполнить и отправить Arasaka Buisness.",
    )
    add_copy_block(doc, [
        "1. Тестовый shopId: ",
        "2. Тестовый Secret key: ",
        "3. Настоящий shopId: ",
        "4. Настоящий Secret key: ",
        "5. Адрес сайта: ",
        f"6. Адрес API-сервера: {API_BASE_URL}",
        "7. Чеки через ЮKassa: да / нет",
        "8. Система налогообложения: ",
        "9. Ставка НДС: ",
        "10. Как проводить курсы в чеке: ",
        "11. Как проводить товары в чеке: ",
        "12. Как проводить пополнение кошелька в чеке: ",
        "13. Подключенные способы оплаты: ",
        "14. Ответственный человек со стороны заказчика: ",
        "15. Контакт для срочных вопросов: ",
    ])

    add_h1(doc, "10. Финальный чек-лист готовности")
    add_checklist_table(doc, [
        ("☐", "Тестовый магазин", "Создан в личном кабинете ЮKassa."),
        ("☐", "Тестовый shopId", "Скопирован и передан Arasaka Buisness."),
        ("☐", "Тестовый Secret key", "Скопирован и передан безопасным личным сообщением."),
        ("☐", "HTTP-уведомления", f"Добавлен URL {WEBHOOK_URL}."),
        ("☐", "События", "Включены payment.succeeded и payment.canceled."),
        ("☐", "Чеки", "Бухгалтер подтвердил, включать их или нет, и дал параметры."),
        ("☐", "Тестовый платеж", "Arasaka Buisness проверил оплату курса/товара."),
        ("☐", "Пополнение кошелька", "Arasaka Buisness проверил пополнение через ЮKassa."),
        ("☐", "Настоящий магазин", "Production shopId и Secret key переданы после успешного теста."),
    ])

    add_h1(doc, "11. Что не делать")
    for text in [
        "Не отправлять Secret key в общий чат.",
        "Не вставлять Secret key на сайт или во фронтенд.",
        "Не включать боевой магазин до успешного тестового платежа.",
        "Не включать чеки без ответа бухгалтера.",
        "Не покупать дополнительные услуги ЮKassa до подтверждения Arasaka Buisness, что они действительно нужны.",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_h1(doc, "12. Что уже подготовлено на стороне сайта")
    for text in [
        "Оплата курсов и товаров через переход на платежную страницу ЮKassa.",
        "Пополнение кошелька через тот же безопасный сценарий.",
        "Получение уведомлений ЮKassa о статусе платежа.",
        "Проверка платежа на сервере перед тем, как отметить заказ оплаченным.",
        "История платежей в личном кабинете пользователя.",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_callout(
        doc,
        "Итог",
        "Со стороны заказчика главное: создать магазины, передать shopId и Secret key, настроить HTTP-уведомления и дать ответ бухгалтера по чекам. "
        "После этого Arasaka Buisness подключит данные на сервере и проведет тестовые платежи.",
        fill="F4F6F9",
    )

    doc.core_properties.title = "Инструкция для подключения ЮKassa"
    doc.core_properties.author = BRAND
    doc.core_properties.subject = f"{PROJECT}: настройка ЮKassa"
    doc.core_properties.comments = "Customer-facing setup instruction prepared by Arasaka Buisness."
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
